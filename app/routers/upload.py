import asyncio
import os
import traceback
import uuid
import shutil
import json
import concurrent.futures

from fastapi import (
    APIRouter,
    Depends,
    UploadFile,
    File,
    HTTPException,
    BackgroundTasks,
    Form,
    Request,
)
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from typing import List, Literal, Optional, Dict, Any
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pydantic import BaseModel, Field
import sqlalchemy
from sqlalchemy.orm import Session

from app.config import Config
from app.database.common import database_context, get_db
from app.database.models import (
    GameDB,
    ProcessingTaskDB,
    ReportDB,
    TeamAnalysisDB,
    TeamDB,
    TeamStatsDB,
    UserDB,
    PlayerDB,
    PlayerStatsDB,
)
from app.llmmodels import GameSimulation, TeamWrapper
from app.routers.util import get_verified_user_email
from app.services.anthropic_api import analyze_team_pdf, simulate_game
from app.services.report_gen import generate_report
from app.database.connection import (
    get_user_by_email,
    insert_team,
    insert_team_stats,
    insert_player,
    insert_player_stats,
    insert_team_analysis,
    insert_game,
    insert_game_simulation,
    insert_report,
    get_recent_analyses,
    execute_query,
    insert_player_raw_stats,
    insert_player_projections,
    insert_simulation_details,
    find_player_by_name,
    update_team_stats_game_id,
)

# Set up Jinja2 templates
config = Config()

router = APIRouter(
    prefix="/task",
    tags=["task"],
    responses={404: {"description": "Not found"}},
)

# Dictionary to store processing status
processing_tasks = {}

# Processing steps for progress tracking
PROCESSING_STEPS = [
    "Analyzing teams statistics",
    "Storing data in database",
    "Generating team analysis report",
    "Simulating game",
    "Generating final report",
    "Your report is ready",
]


class ProcessingTaskResponse(BaseModel):
    task_uuid: str
    status: Literal["processing", "completed", "failed"]
    step_description: str
    current_step: int
    total_steps: int
    game_uuid: Optional[str] = None


class UploadProcessResponse(BaseModel):
    task_id: str
    status: Literal["processing", "completed", "failed"]


@router.get("/analyses", response_class=JSONResponse)
async def get_analyses(request: Request, user_email=Depends(get_verified_user_email)):
    """
    Get recent analyses and display them on a webpage
    """
    user = get_user_by_email(user_email)

    # Get analyses filtered by user_id
    analyses = get_recent_analyses(limit=10, user_id=user["id"])

    return JSONResponse(content={"analyses": analyses})


@router.post("/upload", response_model=UploadProcessResponse)
async def upload_files(
    background_tasks: BackgroundTasks,
    request: Request,
    team_uuid: Optional[str] = Form(None, description="UUID of the team to analyze"),
    team_files: Optional[UploadFile] = File(None, description="PDF file of the team to analyze"),
    team_name: Optional[str] = Form(None, description="Name of the team to analyze"),
    opponent_files: UploadFile = File(..., description="PDF file of the opponent to analyze"),
    opponent_name: Optional[str] = Form(...),
    use_local_simulation: Optional[bool] = Form(False),
    user_email: str = Depends(get_verified_user_email),
    db: Session = Depends(get_db),
):
    """
    Upload PDF files for analysis - one for our team and one for the opponent
    """
    team_db = None
    if team_uuid is not None:
        team_db = db.query(TeamDB).filter(TeamDB.uuid == team_uuid).first()
        if team_db is None:
            raise HTTPException(status_code=404, detail="Team not found")

    # Validate files are PDFs
    if team_uuid is None and not team_files.filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=400, detail=f"Team file {team_files.filename} is not a PDF"
        )

    if not opponent_files.filename.lower().endswith(".pdf"):
        raise HTTPException(
            status_code=400,
            detail=f"Opponent file {opponent_files.filename} is not a PDF",
        )

    # Create a unique task ID
    task_uuid = str(uuid.uuid4())
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    task_dir = f"{config.base_dir}/app/temp/uploads/{task_uuid}"
    os.makedirs(task_dir, exist_ok=True)

    # Save uploaded files
    file_paths = []

    # Save team file
    if team_uuid is None:
        team_file_path = f"{task_dir}/team_{team_files.filename}"
        with open(team_file_path, "wb") as buffer:
            content = await team_files.read()
            buffer.write(content)
        file_paths.append(team_file_path)

    # Save opponent file
    opponent_file_path = f"{task_dir}/opponent_{opponent_files.filename}"
    with open(opponent_file_path, "wb") as buffer:
        content = await opponent_files.read()
        buffer.write(content)
    file_paths.append(opponent_file_path)

    processing_task = ProcessingTaskDB(
        status="processing",
        team_file_path=team_file_path if team_uuid is None else None,
        opponent_file_path=opponent_file_path,
        team_id=team_db.id if team_db else None,
        step=0,
        total_steps=len(PROCESSING_STEPS),
        task_uuid=task_uuid,
    )
    db.add(processing_task)
    db.commit()

    # Get current user from request state
    user = get_user_by_email(db, user_email)

    # Process files in background
    background_tasks.add_task(
        process_files,
        task_uuid,
        user.id,
        team_name,
        opponent_name,
        use_local_simulation,
    )

    return UploadProcessResponse(task_id=task_uuid, status="processing")


@router.get("/status/{task_id}", response_model=ProcessingTaskResponse)
def get_status(task_id: str, db: Session = Depends(get_db)):
    """
    Get the status of a processing task
    """
    processing_task_db = (
        db.query(ProcessingTaskDB).filter(ProcessingTaskDB.task_uuid == task_id).first()
    )
    if processing_task_db is None:
        raise HTTPException(status_code=404, detail="Task not found")

    if processing_task_db.game_id is not None:
        game_db = (
            db.query(GameDB).filter(GameDB.id == processing_task_db.game_id).first()
        )
        game_uuid = str(game_db.uuid)
    else:
        game_uuid = None

    return ProcessingTaskResponse(
        task_uuid=task_id,
        status=processing_task_db.status,
        step_description=PROCESSING_STEPS[processing_task_db.step]
        if processing_task_db.step < len(PROCESSING_STEPS)
        else "Completed",
        current_step=processing_task_db.step,
        total_steps=processing_task_db.total_steps,
        game_uuid=game_uuid,
    )


@router.get("/download/{task_id}")
async def download_report(task_id: str):
    """
    Download the generated report
    """
    # First check if task_id is in processing_tasks
    if task_id in processing_tasks:
        task = processing_tasks[task_id]

        if task["status"] != "completed":
            raise HTTPException(status_code=400, detail="Report not ready yet")

        if not task["report_path"] or not os.path.exists(task["report_path"]):
            raise HTTPException(status_code=404, detail="Report file not found")

        return FileResponse(
            path=task["report_path"],
            filename=os.path.basename(task["report_path"]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        # If not in processing_tasks, check if it's a report ID in the database
        query = """
        SELECT r.file_path
        FROM reports r
        WHERE r.id = %s
        """

        result = execute_query(query, (task_id,))

        if (
            not result
            or not result[0]["file_path"]
            or not os.path.exists(result[0]["file_path"])
        ):
            raise HTTPException(status_code=404, detail="Report not found")

        return FileResponse(
            path=result[0]["file_path"],
            filename=os.path.basename(result[0]["file_path"]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


@router.get("/download-team-analysis/{task_id}")
async def download_team_analysis(task_id: str):
    """
    Download the team analysis report
    """
    # First check if task_id is in processing_tasks
    if task_id in processing_tasks:
        task = processing_tasks[task_id]

        if task["status"] != "completed":
            raise HTTPException(status_code=400, detail="Reports not ready yet")

        if not task.get("team_analysis_path") or not os.path.exists(
            task.get("team_analysis_path", "")
        ):
            raise HTTPException(
                status_code=404, detail="Team analysis report not found"
            )

        return FileResponse(
            path=task["team_analysis_path"],
            filename=os.path.basename(task["team_analysis_path"]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        # If not in processing_tasks, check if it's a report ID in the database
        query = """
        SELECT r.file_path
        FROM reports r
        JOIN games g ON r.game_id = g.id
        WHERE r.id = %s AND r.report_type = 'team_analysis'
        """

        result = execute_query(query, (task_id,))

        if (
            not result
            or not result[0]["file_path"]
            or not os.path.exists(result[0]["file_path"])
        ):
            raise HTTPException(
                status_code=404, detail="Team analysis report not found"
            )

        return FileResponse(
            path=result[0]["file_path"],
            filename=os.path.basename(result[0]["file_path"]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


@router.get("/download-opponent-analysis/{task_id}")
async def download_opponent_analysis(task_id: str):
    """
    Download the opponent analysis report
    """
    # First check if task_id is in processing_tasks
    if task_id in processing_tasks:
        task = processing_tasks[task_id]

        if task["status"] != "completed":
            raise HTTPException(status_code=400, detail="Reports not ready yet")

        if not task.get("opponent_analysis_path") or not os.path.exists(
            task.get("opponent_analysis_path", "")
        ):
            raise HTTPException(
                status_code=404, detail="Opponent analysis report not found"
            )

        return FileResponse(
            path=task["opponent_analysis_path"],
            filename=os.path.basename(task["opponent_analysis_path"]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    else:
        # If not in processing_tasks, check if it's a report ID in the database
        query = """
        SELECT r.file_path
        FROM reports r
        JOIN games g ON r.game_id = g.id
        WHERE r.id = %s AND r.report_type = 'opponent_analysis'
        """

        result = execute_query(query, (task_id,))

        if (
            not result
            or not result[0]["file_path"]
            or not os.path.exists(result[0]["file_path"])
        ):
            raise HTTPException(
                status_code=404, detail="Opponent analysis report not found"
            )

        return FileResponse(
            path=result[0]["file_path"],
            filename=os.path.basename(result[0]["file_path"]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


@router.get("/download-by-path")
async def download_by_path(path: str):
    """
    Download a report by its file path
    """
    if not path or not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Report file not found")

    return FileResponse(
        path=path,
        filename=os.path.basename(path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def generate_team_analysis_report(db, team_id: int) -> str:
    """
    Generate a DOCX report for team analysis

    Args:
        db: Database session
        team_analysis_id: ID of the team analysis to generate report for

    Returns:
        Path to the generated report
    """
    # Get the team
    team = db.query(TeamDB).filter(TeamDB.id == team_id).first()
    if not team:
        raise ValueError(f"No team found with ID {team_id}")

    # Get the latest team analysis and stats
    team_analysis = (
        db.query(TeamAnalysisDB)
        .filter(TeamAnalysisDB.team_id == team_id)
        .order_by(TeamAnalysisDB.id.desc())
        .first()
    )
    
    team_stats = (
        db.query(TeamStatsDB)
        .filter(TeamStatsDB.team_id == team_id)
        .order_by(TeamStatsDB.id.desc())
        .first()
    )

    if not team_analysis or not team_stats:
        raise ValueError(f"No analysis or stats found for team with ID {team_id}")

    # Create a new document
    doc = Document()

    # Add title
    title = doc.add_heading(f"{team.name} - Team Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add team info section
    doc.add_heading("Team Information", 1)
    doc.add_paragraph(f"Team: {team.name}")
    doc.add_paragraph(f"Record: {team.record or 'N/A'}")
    if team.ranking:
        doc.add_paragraph(f"Ranking: {team.ranking}")

    # Add team stats section
    doc.add_heading("Team Statistics", 1)
    stats_table = doc.add_table(rows=1, cols=2)
    stats_table.style = "Table Grid"

    # Add header row
    header_cells = stats_table.rows[0].cells
    header_cells[0].text = "Statistic"
    header_cells[1].text = "Value"

    # Add stats rows
    stats_dict = {
        "Points Per Game": team_stats.ppg,
        "Field Goal %": team_stats.fg_pct,
        "3-Point %": team_stats.fg3_pct,
        "Free Throw %": team_stats.ft_pct,
        "Rebounds": team_stats.rebounds,
        "Offensive Rebounds": team_stats.offensive_rebounds,
        "Defensive Rebounds": team_stats.defensive_rebounds,
        "Assists": team_stats.assists,
        "Steals": team_stats.steals,
        "Blocks": team_stats.blocks,
        "Turnovers": team_stats.turnovers,
        "Assist/Turnover": team_stats.assist_to_turnover,
    }

    for stat, value in stats_dict.items():
        row_cells = stats_table.add_row().cells
        row_cells[0].text = stat
        row_cells[1].text = str(value)

    # Add team strengths section
    doc.add_heading("Team Strengths", 1)
    for strength in team_analysis.strengths:
        doc.add_paragraph(f"• {strength}", style="List Bullet")

    # Add team weaknesses section
    doc.add_heading("Team Weaknesses", 1)
    for weakness in team_analysis.weaknesses:
        doc.add_paragraph(f"• {weakness}", style="List Bullet")

    # Add key players section
    doc.add_heading("Key Players", 1)
    for player in team_analysis.key_players:
        doc.add_paragraph(f"• {player}", style="List Bullet")

    # Add playing style section
    if team_analysis.playing_style:
        doc.add_heading("Playing Style", 1)
        doc.add_paragraph(team_analysis.playing_style)

    # Add offensive keys section
    doc.add_heading("Offensive Keys", 1)
    for key in team_analysis.offensive_keys:
        doc.add_paragraph(f"• {key}", style="List Bullet")

    # Add defensive keys section
    doc.add_heading("Defensive Keys", 1)
    for key in team_analysis.defensive_keys:
        doc.add_paragraph(f"• {key}", style="List Bullet")

    # Add game factors section
    doc.add_heading("Game Factors", 1)
    for factor in team_analysis.game_factors:
        doc.add_paragraph(f"• {factor}", style="List Bullet")

    # Add rotation plan section
    if team_analysis.rotation_plan:
        doc.add_heading("Rotation Plan", 1)
        for plan in team_analysis.rotation_plan:
            doc.add_paragraph(f"• {plan}", style="List Bullet")

    # Add situational adjustments section
    doc.add_heading("Situational Adjustments", 1)
    for adjustment in team_analysis.situational_adjustments:
        doc.add_paragraph(f"• {adjustment}", style="List Bullet")

    # Add game keys section
    doc.add_heading("Game Keys", 1)
    for key in team_analysis.game_keys:
        doc.add_paragraph(f"• {key}", style="List Bullet")

    # Add player details section
    doc.add_heading("Player Details", 1)
    players = db.query(PlayerDB).filter(PlayerDB.team_id == team.id).limit(5).all()

    for player in players:
        doc.add_heading(f"{player.name} #{player.number}", 2)

        # Add player stats
        player_stats = (
            db.query(PlayerStatsDB)
            .filter(
                PlayerStatsDB.player_id == player.id,
                PlayerStatsDB.is_season_average.is_(True),
            )
            .first()
        )

        if player_stats:
            player_stats_table = doc.add_table(rows=1, cols=2)
            player_stats_table.style = "Table Grid"

            # Add header row
            header_cells = player_stats_table.rows[0].cells
            header_cells[0].text = "Statistic"
            header_cells[1].text = "Value"

            # Add stats rows
            stats_dict = {
                "Games Played": player_stats.games_played,
                "Points Per Game": player_stats.ppg,
                "Field Goal %": player_stats.fg_pct,
                "3-Point %": player_stats.fg3_pct,
                "Free Throw %": player_stats.ft_pct,
                "Rebounds": player_stats.rpg,
                "Assists": player_stats.apg,
                "Steals": player_stats.spg,
                "Blocks": player_stats.bpg,
                "Turnovers": player_stats.topg,
                "Minutes": player_stats.minutes,
            }

            for stat, value in stats_dict.items():
                row_cells = player_stats_table.add_row().cells
                row_cells[0].text = stat
                row_cells[1].text = str(value)

        # Add player strengths
        if player.strengths:
            doc.add_heading("Strengths", 3)
            for strength in player.strengths:
                doc.add_paragraph(f"• {strength}", style="List Bullet")

        # Add player weaknesses
        if player.weaknesses:
            doc.add_heading("Weaknesses", 3)
            for weakness in player.weaknesses:
                doc.add_paragraph(f"• {weakness}", style="List Bullet")

    # Save the document
    report_dir = f"{config.base_dir}/app/temp/reports"
    os.makedirs(report_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_filename = f"{team.name} - Team Analysis - {timestamp}.docx"
    report_path = os.path.join(report_dir, report_filename)
    doc.save(report_path)

    return report_path


def run_team_analysis(
    task_uuid: str, file_path: str, is_home_team: bool, team_name: str
):
    """
    Run the team analysis
    """
    with database_context() as db:
        query = (
            sqlalchemy.update(ProcessingTaskDB)
            .where(ProcessingTaskDB.task_uuid == task_uuid)
            .values(step=0, status="processing")
        )
        db.execute(query)
        db.commit()

        print(f"DEBUG - Starting team analysis for {team_name} with file path {file_path}")
        # team_wrapper = analyze_team_pdf(file_path, is_our_team=is_home_team)
        # with open("/Users/edoardo/programming/anova/team2_wrapper.json", "w+") as f:
        #     json.dump(opponent_wrapper.model_dump(mode="json"), f)
        with open(f"/Users/edoardo/programming/anova/team{1 if is_home_team else 2}_wrapper.json", "r") as f:
            team_wrapper = TeamWrapper.model_validate(json.load(f))

        team_label = "home" if is_home_team else "away"
        team_analysis = team_wrapper.team_analysis
        # print("-"*40 + "\n" + "DEBUG - Opponent Analysis:", opponent_analysis)

        # Override team names if provided
        if team_name:
            team_wrapper.team_details.team_name = team_name

        query = (
            sqlalchemy.update(ProcessingTaskDB)
            .where(ProcessingTaskDB.task_uuid == task_uuid)
            .values(step=1, status="processing")
        )
        db.execute(query)
        db.commit()

        # Insert teams into database
        print(f"DEBUG - Inserting {team_label} team into database")
        team_id = insert_team(db, team_wrapper.team_details)

        print(f"DEBUG - {team_label} team ID: {team_id}")

        # Insert team stats
        print(f"DEBUG - Inserting {team_label} stats into database")
        team_stats_id = insert_team_stats(db, team_id, team_wrapper.team_stats)
        print(f"DEBUG - {team_label} Stats ID: {team_stats_id}")

        print(f"DEBUG - Inserting {team_label} players and their stats into database")
        for player in team_wrapper.team_details.players:
            player_id = insert_player(db, team_id, player)
            print(f"DEBUG - {team_label} Player ID: {player_id}, Name: {player.name}")
            if player_id:
                # Insert raw stats first
                raw_stats_id = insert_player_raw_stats(db, player_id, player.stats)
                print(f"DEBUG - {team_label} Player Raw Stats ID: {raw_stats_id}")
                # Then insert processed stats with reference to raw stats
                player_stats_id = insert_player_stats(
                    db, player_id, player.stats, player_raw_stats_id=raw_stats_id
                )
                print(f"DEBUG - {team_label} Player Stats ID: {player_stats_id}")

        print(f"DEBUG - Inserting {team_label} analysis into database")
        team_analysis_id = insert_team_analysis(db, team_id, team_analysis)
        print(f"DEBUG - {team_label} Analysis ID: {team_analysis_id}")

        return team_id, team_stats_id, team_analysis_id


# it must NOT be async, otherwise it will mess with the fast api event loop and continuously
# freeze it
def process_files(
    task_uuid: str,
    user_id: int,
    team_name: Optional[str],
    opponent_name: Optional[str],
    use_local_simulation: bool = False,
):
    """
    Process uploaded PDF files and generate a report
    """
    with database_context() as db:
        user = db.query(UserDB).filter(UserDB.id == user_id).first()
        processing_task_db = (
            db.query(ProcessingTaskDB)
            .filter(ProcessingTaskDB.task_uuid == task_uuid)
            .first()
        )
        if not user:
            raise ValueError("User not found")

        try:
            # Find team and opponent file paths
            team_id = processing_task_db.team_id
            team_uuid = str(db.query(TeamDB.uuid).where(TeamDB.id == team_id).first().uuid)
            team_file_path = processing_task_db.team_file_path
            opponent_file_path = processing_task_db.opponent_file_path

            if team_uuid is None and team_file_path is None:
                sql_query = (
                    sqlalchemy.update(ProcessingTaskDB)
                    .where(ProcessingTaskDB.task_uuid == task_uuid)
                    .values(team_uuid=team_uuid)
                )
                db.execute(sql_query)
                db.commit()

            if opponent_file_path is None:
                raise ValueError("Could not identify opponent file")

            if (
                processing_task_db.team_file_path is None
                and processing_task_db.opponent_file_path is None
            ) or not opponent_file_path:
                raise ValueError("Could not identify team and opponent files")

            print("-" * 40 + "\n" + f"DEBUG - Processing team file: {team_file_path}")
            print(f"DEBUG - Processing opponent file: {opponent_file_path}")

            if team_uuid is None:
                # If both team are provided as files, we do parallel analysis
                with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
                    team_future = executor.submit(
                        run_team_analysis, task_uuid, team_file_path, True, team_name
                    )
                    opponent_future = executor.submit(
                        run_team_analysis,
                        task_uuid,
                        opponent_file_path,
                        False,
                        opponent_name,
                    )
                    
                    try:
                        team_result = team_future.result()
                        opponent_result = opponent_future.result()
                    except Exception as e:
                        print(f"DEBUG - Error in team analysis: {e}")
                        query = (
                            sqlalchemy.update(ProcessingTaskDB)
                            .where(ProcessingTaskDB.task_uuid == task_uuid)
                            .values(step=0, status="failed")
                        )
                        db.execute(query)
                        db.commit()
                        raise e


                    team_id, team_stats_id, team_analysis_id = team_result
                    opponent_id, opponent_stats_id, opponent_analysis_id = (
                        opponent_result
                    )
            else:
                # fetch already existing stats for team, and do analysis for opponent
                team_db = db.query(TeamDB.id).where(TeamDB.uuid == team_uuid).first()
                if team_db is None:
                    query = (
                        sqlalchemy.update(ProcessingTaskDB)
                        .where(ProcessingTaskDB.task_uuid == task_uuid)
                        .values(step=0, status="failed")
                    )
                    db.execute(query)
                    db.commit()
                    raise ValueError("Team doesn't exist")
                
                team_id = team_db.id
                team_analysis_db = (
                    db.query(TeamAnalysisDB.id)
                    .where(TeamAnalysisDB.team_id == team_id)
                    .order_by(TeamAnalysisDB.id.desc())
                    .first()
                )
                team_stats_db = (
                    db.query(TeamStatsDB)
                    .where(TeamStatsDB.team_id == team_id)
                    .order_by(TeamStatsDB.id.desc())
                    .first()
                )

                if team_analysis_db is None or team_stats_db is None:
                    raise ValueError("Analysis or stats don't exist")
                
                team_analysis_id, team_stats_id = team_analysis_db.id, team_stats_db.id

                try:
                    opponent_id, opponent_stats_id, opponent_analysis_id = run_team_analysis(
                        task_uuid, opponent_file_path, False, opponent_name
                    )
                except Exception as e:
                    print(f"DEBUG - Error in opponent analysis: {e}")
                    query = (
                        sqlalchemy.update(ProcessingTaskDB)
                        .where(ProcessingTaskDB.task_uuid == task_uuid)
                        .values(step=0, status="failed")
                    )
                    raise e

            # Insert game with user ID if available
            print("DEBUG - Inserting game into database")
            user_id = user.id

            game_id, game_uuid = insert_game(db, team_id, opponent_id, user_id)
            print(f"DEBUG - Game ID: {game_id}, Game UUID: {game_uuid}")

            # Set the game id for the team and opponent stats
            update_team_stats_game_id(db, team_stats_id, game_id)
            update_team_stats_game_id(db, opponent_stats_id, game_id)

            # Step 5: Generate opponent analysis report
            query = (
                sqlalchemy.update(ProcessingTaskDB)
                .where(ProcessingTaskDB.task_uuid == task_uuid)
                .values(step=2, status="processing")
            )
            db.execute(query)
            db.commit()

            team_analysis_path = generate_team_analysis_report(db, team_id)
            opponent_analysis_path = generate_team_analysis_report(
                db, opponent_id
            )

            # Insert reports
            if game_id:
                print("DEBUG - Inserting reports into database")
                team_report_id, _ = insert_report(
                    db, game_id, "team_analysis", team_analysis_path
                )
                opponent_report_id, _ = insert_report(
                    db, game_id, "opponent_analysis", opponent_analysis_path
                )
                print(
                    f"DEBUG - Team Report ID: {team_report_id}, Opponent Report ID: {opponent_report_id}"
                )

            # Step 6: Simulate game
            query = (
                sqlalchemy.update(ProcessingTaskDB)
                .where(ProcessingTaskDB.task_uuid == task_uuid)
                .values(step=3, status="processing")
            )
            db.execute(query)
            db.commit()

            # with open("/Users/edoardo/programming/anova/simulation_results.json", "r") as f:
            #     simulation_results = GameSimulation.model_validate(json.load(f))
            simulation_results = simulate_game(
                db, team_id, opponent_id, use_local=use_local_simulation
            )
            # with open("/Users/edoardo/programming/anova/simulation_results.json", "w+") as f:
            #     json.dump(simulation_results.model_dump(mode="json"), f)

            simulation_results_dict = simulation_results.model_dump(mode="json")
            # print("-"*40 + "\n" + "DEBUG - Simulation Results:", simulation_results)

            # Insert game simulation
            if game_id:
                print("DEBUG - Inserting game simulation into database")
                simulation_id = insert_game_simulation(db, game_id, simulation_results)
                print(f"DEBUG - Simulation ID: {simulation_id}")

                # If using local simulation, insert simulation details
                if use_local_simulation and "numSimulations" in simulation_results:
                    print("DEBUG - Inserting simulation details into database")
                    simulation_details_id = insert_simulation_details(
                        db,
                        simulation_id,
                        game_id,
                        team_id,  # home team
                        opponent_id,  # away team
                        simulation_results_dict,
                    )
                    print(f"DEBUG - Simulation Details ID: {simulation_details_id}")

                # Insert player projections
                print("DEBUG - Inserting player projections into database")

                # Process team player projections
                for i in range(1, 7):  # Assuming up to 6 players
                    player_key = f"team_p{i}_name"
                    if player_key in simulation_results_dict:
                        # Find player ID by name
                        player_name = simulation_results_dict[f"team_p{i}_name"]
                        player = find_player_by_name(db, player_name, team_id)

                        if player:
                            projection_data = {
                                "ppg": simulation_results_dict.get(f"team_p{i}_ppg", 0),
                                "rpg": simulation_results_dict.get(f"team_p{i}_rpg", 0),
                                "apg": simulation_results_dict.get(f"team_p{i}_apg", 0),
                                "fg": simulation_results_dict.get(
                                    f"team_p{i}_fg", "0%"
                                ),
                                "3p": simulation_results_dict.get(
                                    f"team_p{i}_3p", "0%"
                                ),
                                "role": simulation_results_dict.get(
                                    f"team_p{i}_role", ""
                                ),
                            }

                            projection_id = insert_player_projections(
                                db,
                                simulation_id,
                                player.id,
                                team_id,
                                game_id,
                                projection_data,
                                True,  # is_home_team
                            )
                            print(
                                f"DEBUG - Team Player Projection ID: {projection_id}, Player: {player_name}"
                            )

                # Process opponent player projections
                for i in range(1, 7):  # Assuming up to 6 players
                    player_key = f"opp_p{i}_name"
                    if player_key in simulation_results_dict:
                        # Find player ID by name
                        player_name = simulation_results_dict[f"opp_p{i}_name"]
                        player = find_player_by_name(db, player_name, opponent_id)

                        if player:
                            projection_data = {
                                "ppg": simulation_results_dict.get(f"opp_p{i}_ppg", 0),
                                "rpg": simulation_results_dict.get(f"opp_p{i}_rpg", 0),
                                "apg": simulation_results_dict.get(f"opp_p{i}_apg", 0),
                                "fg": simulation_results_dict.get(f"opp_p{i}_fg", "0%"),
                                "3p": simulation_results_dict.get(f"opp_p{i}_3p", "0%"),
                                "role": simulation_results_dict.get(
                                    f"opp_p{i}_role", ""
                                ),
                            }

                            projection_id = insert_player_projections(
                                db,
                                simulation_id,
                                player.id,
                                opponent_id,
                                game_id,
                                projection_data,
                                False,  # is_home_team
                            )
                            print(
                                f"DEBUG - Opponent Player Projection ID: {projection_id}, Player: {player_name}"
                            )

            # Step 7: Generate final report
            print("DEBUG - Generating final report")
            query = (
                sqlalchemy.update(ProcessingTaskDB)
                .where(ProcessingTaskDB.task_uuid == task_uuid)
                .values(step=4, status="processing")
            )
            db.execute(query)
            db.commit()

            report_path = generate_report(db, game_id)
            report_id, _ = insert_report(db, game_id, "game_analysis", report_path)
            print(f"DEBUG - Report ID: {report_id}")

            query = (
                sqlalchemy.update(ProcessingTaskDB)
                .where(ProcessingTaskDB.task_uuid == task_uuid)
                .values(step=5, status="completed", game_id=game_id)
            )
            db.execute(query)
            db.commit()

        except Exception as e:
            # Update task status with error
            query = (
                sqlalchemy.update(ProcessingTaskDB)
                .where(ProcessingTaskDB.task_uuid == task_uuid)
                .values(status="failed")
            )
            db.execute(query)
            print("-" * 40)
            print(f"ERROR: {str(e)} : {traceback.format_exc()}")


if __name__ == "__main__":
    processing_tasks["0d05182f-2088-418f-bd4e-fed202f8a271"] = {}
    process_files(
        "0d05182f-2088-418f-bd4e-fed202f8a271",
        1,
        "ARLINGTON",
        "SCARSDALE",
        False,
    )
