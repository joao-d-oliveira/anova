import os
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from sqlalchemy.orm import Session

from app.config import Config
from app.database.models import (
    TeamDB, 
    GameDB, 
    TeamAnalysisDB, 
    GameSimulationDB,
    PlayerProjectionDB,
    TeamStatsDB,
    PlayerStatsDB
)

config = Config()

def generate_report(db: Session, game_id: int) -> str:
    """
    Generate a DOCX report based on game data from the database
    
    Args:
        db: SQLAlchemy database session
        game_id: ID of the game to generate report for
        
    Returns:
        Path to the generated report
    """
    # Fetch game data
    game = db.query(GameDB).filter(GameDB.id == game_id).first()
    if not game:
        raise ValueError(f"Game with ID {game_id} not found")
    
    # Fetch team data
    home_team = db.query(TeamDB).filter(TeamDB.id == game.home_team_id).first()
    away_team = db.query(TeamDB).filter(TeamDB.id == game.away_team_id).first()
    
    # Fetch team analysis
    home_analysis = db.query(TeamAnalysisDB).filter(TeamAnalysisDB.team_id == home_team.id).first()
    away_analysis = db.query(TeamAnalysisDB).filter(TeamAnalysisDB.team_id == away_team.id).first()
    
    # Fetch team stats
    home_stats = db.query(TeamStatsDB).filter(
        TeamStatsDB.team_id == home_team.id,
        TeamStatsDB.is_season_average == True
    ).first()
    
    away_stats = db.query(TeamStatsDB).filter(
        TeamStatsDB.team_id == away_team.id,
        TeamStatsDB.is_season_average == True
    ).first()
    
    # Fetch game simulation
    simulation = db.query(GameSimulationDB).filter(GameSimulationDB.game_id == game_id).first()
    
    # Create a new Document
    doc = Document()
    
    # Set document properties
    doc.core_properties.title = f"{home_team.name} VS {away_team.name} - Anova Analysis"
    doc.core_properties.author = "Anova Basketball Analytics"
    
    # Add title
    add_title(doc, home_team.name, away_team.name)
    
    # Add matchup overview section
    add_matchup_overview(doc, game, simulation)
    
    # Add team comparison section
    add_team_comparison(doc, home_team, away_team, home_stats, away_stats, home_analysis, away_analysis)
    
    # Add game plan section
    add_game_plan(doc, home_analysis)
    
    # Add team analysis section
    add_team_analysis(doc, home_team, home_analysis)
    
    # Add opponent analysis section
    add_opponent_analysis(doc, away_team, away_analysis)
    
    # Add simulation results section
    add_simulation_results(doc, simulation)
    
    # Create output directory if it doesn't exist
    os.makedirs(f"{config.base_dir}/app/temp/reports", exist_ok=True)
    
    # Generate filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{home_team.name} VS {away_team.name} - Anova Analysis - {timestamp}.docx"
    filepath = os.path.join(f"{config.base_dir}/app/temp/reports", filename)
    
    # Save the document
    doc.save(filepath)
    
    return filepath

def add_title(doc: Document, home_team_name: str, away_team_name: str) -> None:
    """Add title section to the document"""
    # Add title
    title = doc.add_heading(f"{home_team_name} VS {away_team_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph("Anova Basketball Analytics")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style = "Subtitle"
    
    # Add date
    date_paragraph = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add horizontal line
    doc.add_paragraph("_" * 50)
    
    # Add some space
    doc.add_paragraph()

def add_matchup_overview(doc: Document, game: GameDB, simulation: Optional[GameSimulationDB]) -> None:
    """Add matchup overview section to the document"""
    # Add section heading
    doc.add_heading("MATCHUP OVERVIEW", level=1)
    
    # Add matchup overview
    overview = doc.add_paragraph()
    overview.add_run(f"Game between {game.home_team.name} and {game.away_team.name} at {game.location}").bold = True
    
    # Add win probability and projected score
    if simulation:
        if simulation.win_probability:
            doc.add_paragraph(simulation.win_probability)
        if simulation.projected_score:
            projected_score = doc.add_paragraph()
            projected_score.add_run("Projected Score: ").bold = True
            projected_score.add_run(simulation.projected_score)
    
    # Add horizontal line
    doc.add_paragraph("_" * 50)

def add_team_comparison(
    doc: Document, 
    home_team: TeamDB, 
    away_team: TeamDB,
    home_stats: TeamStatsDB,
    away_stats: TeamStatsDB,
    home_analysis: TeamAnalysisDB,
    away_analysis: TeamAnalysisDB
) -> None:
    """Add team comparison section to the document"""
    # Add section heading
    doc.add_heading("TEAM COMPARISON", level=1)
    
    # Create a table for team comparison
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "STATISTIC"
    header_cells[1].text = home_team.name
    header_cells[2].text = away_team.name
    
    # Format header row
    for cell in header_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    # Add stats rows
    stats = [
        ("PPG", home_stats.ppg, away_stats.ppg),
        ("FG%", home_stats.fg_pct, away_stats.fg_pct),
        ("3P%", home_stats.fg3_pct, away_stats.fg3_pct),
        ("FT%", home_stats.ft_pct, away_stats.ft_pct),
        ("RPG", home_stats.rebounds, away_stats.rebounds),
        ("APG", home_stats.assists, away_stats.assists),
        ("SPG", home_stats.steals, away_stats.steals),
        ("BPG", home_stats.blocks, away_stats.blocks),
        ("TOPG", home_stats.turnovers, away_stats.turnovers),
        ("A/TO", home_stats.assist_to_turnover, away_stats.assist_to_turnover)
    ]
    
    for label, home_stat, away_stat in stats:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(home_stat)
        row_cells[2].text = str(away_stat)
        
        # Center align the cells
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space
    doc.add_paragraph()
    
    # Add team strengths and weaknesses
    doc.add_heading(f"{home_team.name} Strengths", level=2)
    add_bullet_list(doc, home_analysis.strengths)
    
    doc.add_heading(f"{home_team.name} Weaknesses", level=2)
    add_bullet_list(doc, home_analysis.weaknesses)
    
    doc.add_heading(f"{away_team.name} Strengths", level=2)
    add_bullet_list(doc, away_analysis.strengths)
    
    doc.add_heading(f"{away_team.name} Weaknesses", level=2)
    add_bullet_list(doc, away_analysis.weaknesses)
    
    # Add horizontal line
    doc.add_paragraph("_" * 50)

def add_game_plan(doc: Document, team_analysis: TeamAnalysisDB) -> None:
    """Add game plan section to the document"""
    # Add section heading
    doc.add_heading("GAME PLAN", level=1)
    
    # Add game factors
    doc.add_heading("Game Factors", level=2)
    add_bullet_list(doc, team_analysis.game_factors)
    
    # Add offensive keys
    doc.add_heading("Offensive Keys", level=2)
    add_bullet_list(doc, team_analysis.offensive_keys)
    
    # Add defensive keys
    doc.add_heading("Defensive Keys", level=2)
    add_bullet_list(doc, team_analysis.defensive_keys)
    
    # Add rotation plan
    doc.add_heading("Rotation Plan", level=2)
    add_bullet_list(doc, team_analysis.rotation_plan)
    
    # Add situational adjustments
    doc.add_heading("Situational Adjustments", level=2)
    add_bullet_list(doc, team_analysis.situational_adjustments)
    
    # Add game keys
    doc.add_heading("Game Keys", level=2)
    add_bullet_list(doc, team_analysis.game_keys)
    
    # Add horizontal line
    doc.add_paragraph("_" * 50)

def add_team_analysis(doc: Document, team: TeamDB, team_analysis: TeamAnalysisDB) -> None:
    """Add team analysis section to the document"""
    # Add section heading
    doc.add_heading(f"{team.name} PLAYER ANALYSIS", level=1)
    
    # Add key players
    doc.add_heading("Key Players", level=2)
    add_bullet_list(doc, team_analysis.key_players)
    
    # Add playing style
    if team_analysis.playing_style:
        style = doc.add_paragraph()
        style.add_run("Playing Style: ").bold = True
        style.add_run(team_analysis.playing_style)
    
    # Add horizontal line
    doc.add_paragraph("_" * 50)

def add_opponent_analysis(doc: Document, team: TeamDB, team_analysis: TeamAnalysisDB) -> None:
    """Add opponent analysis section to the document"""
    # Add section heading
    doc.add_heading(f"{team.name} PLAYER ANALYSIS", level=1)
    
    # Add key players
    doc.add_heading("Key Players", level=2)
    add_bullet_list(doc, team_analysis.key_players)
    
    # Add playing style
    if team_analysis.playing_style:
        style = doc.add_paragraph()
        style.add_run("Playing Style: ").bold = True
        style.add_run(team_analysis.playing_style)
    
    # Add horizontal line
    doc.add_paragraph("_" * 50)

def add_simulation_results(doc: Document, simulation: Optional[GameSimulationDB]) -> None:
    """Add simulation results section to the document"""
    if not simulation:
        return
        
    # Add section heading
    doc.add_heading("SIMULATION RESULTS", level=1)
    
    # Add simulation overview
    if simulation.sim_overall_summary:
        overview = doc.add_paragraph()
        overview.add_run("Overview: ").bold = True
        overview.add_run(simulation.sim_overall_summary)
    
    # Add success factors
    if simulation.sim_success_factors:
        doc.add_heading("Success Factors", level=2)
        add_bullet_list(doc, simulation.sim_success_factors.replace("-", "").split("\n"))
    
    # Add key matchups
    if simulation.sim_key_matchups:
        doc.add_heading("Key Matchups", level=2)
        add_bullet_list(doc, simulation.sim_key_matchups.replace("-", "").split("\n"))
    
    # Add win/loss patterns
    if simulation.sim_win_loss_patterns:
        doc.add_heading("Win/Loss Patterns", level=2)
        add_bullet_list(doc, simulation.sim_win_loss_patterns.replace("-", "").split("\n"))
    
    # Add critical advantage
    if simulation.sim_critical_advantage:
        doc.add_heading("Critical Advantage", level=2)
        add_bullet_list(doc, simulation.sim_critical_advantage.replace("-", "").split("\n"))
    
    # Add keys to victory
    if simulation.sim_keys_to_victory:
        doc.add_heading("Keys to Victory", level=2)
        add_bullet_list(doc, simulation.sim_keys_to_victory)

def add_bullet_list(doc: Document, items: list[str]) -> None:
    """Add a bullet list to the document"""
    if not items:
        return
    
    for item in items:
        if item:
            p = doc.add_paragraph(item, style="List Bullet")
